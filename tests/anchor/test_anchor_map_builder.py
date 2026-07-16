#!/usr/bin/env python3
"""
RED test -- anchor-map builder (issue #75).

Issue #3 built the section-anchor map artifact and its two governance gates
(heading-hash drift, form coverage). Issue #75 is the builder's own contract:
`scripts/build_anchor_map.py` is the artifact everything in the detector layer
depends on, so the builder itself needs fixture coverage independent of the
downstream gates:

  1. Synthetic-mode fixture: the builder's SECTION_CONFIG produces the expected
     anchor set, including the hand-configured sec-10 sub-clause splits
     (notices/exclusivity/merger/precedence) with `sub_clause_split: true` and
     the correct `parent_section`.
  2. Heading-hash stability: the SAME section config run twice produces the
     SAME heading_hash per anchor (same input -> same hashes).
  3. Real-.docx mode: a synthetic .docx fixture built with python-docx, fed
     through `build_anchors_from_docx()`, must resolve every SECTION_CONFIG
     anchor to a heading actually present in the .docx (no silent fallback to
     the config heading -- that would mask real drift between the config and
     the bundled form). Like `build_anchors_from_docx()` itself, this gate
     treats python-docx as optional -- consistent with every other gate in
     this repo ("no pip install needed, tests use only the stdlib"): if
     python-docx is not installed, this gate is SKIPPED (not failed), so CI
     does not gain a new hard dependency. Installing python-docx locally
     exercises the real assertion.
  4. Determinism (byte-identical artifact): running the builder twice on the
     same input must produce a BYTE-IDENTICAL output file, not merely an
     identical `anchor_map_hash`. ARCHITECTURE.md's determinism convention
     (see scripts/diff_standard_form.py: serialize_diff()/diff_hash() have no
     timestamp in the hashed content) applies equally here: the anchor map is
     a release-bundle artifact and "same input -> byte-identical map" is the
     issue #75 acceptance criterion verbatim. A `generated_at` wall-clock
     timestamp embedded in the written file breaks this.

This test FAILS today (RED) because:
  - build_anchor_map.py's main() embeds `generated_at` (datetime.now()) in the
    JSON file it writes, so two runs on identical input never produce
    byte-identical files (only the anchor_map_hash / standard_form_hash
    sub-fields are stable).

Exit codes: 0 = pass, 1 = fail
"""

import json
import subprocess
import sys
import tempfile
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
SCRIPTS_DIR = REPO_ROOT / "scripts"

sys.path.insert(0, str(SCRIPTS_DIR))


def _import_builder():
    try:
        import build_anchor_map  # type: ignore
        return build_anchor_map, ""
    except ImportError as exc:
        return None, f"MISSING: scripts/build_anchor_map.py does not import ({exc})."


def check_section_config_fixture(mod):
    """Gate 1: synthetic SECTION_CONFIG produces the expected anchor set,
    including the sec-10 sub-clause splits.

    Explicitly resolved for "eiaa" (issue #343 repointed the registry
    default to the public "sample-agreement" sample playbook, whose
    deliberately small section config has no sec-10 sub-clause splits at
    all) -- this gate's whole point is the sec-10 hand-split config, which
    only eiaa's section-config data file carries."""
    failures = []
    eiaa_section_config = mod.load_section_config("eiaa")
    anchors = mod.build_anchors_from_config(
        config=eiaa_section_config["sections"],
        absent_from_form_anchors=eiaa_section_config["absent_from_form_anchors"],
        structural_anchors=eiaa_section_config["structural_anchors"],
    )

    expected_sub_splits = {
        "sec-10-notices": "sec-10",
        "sec-10-non-exclusive": "sec-10",
        "sec-10-merger": "sec-10",
        "sec-10-precedence": "sec-10",
    }

    for anchor, parent in expected_sub_splits.items():
        entry = anchors.get(anchor)
        if entry is None:
            failures.append(f"[G1] MISSING anchor '{anchor}' in synthetic SECTION_CONFIG output.")
            continue
        if entry.get("sub_clause_split") is not True:
            failures.append(
                f"[G1] '{anchor}' must have sub_clause_split: true (got {entry.get('sub_clause_split')!r})."
            )
        if entry.get("parent_section") != parent:
            failures.append(
                f"[G1] '{anchor}' parent_section must be '{parent}' (got {entry.get('parent_section')!r})."
            )

    # Non-split anchors must NOT carry a parent_section.
    non_split_sample = "sec-8"
    entry = anchors.get(non_split_sample)
    if entry is None:
        failures.append(f"[G1] MISSING expected non-split anchor '{non_split_sample}'.")
    elif entry.get("sub_clause_split") is not False:
        failures.append(f"[G1] '{non_split_sample}' must have sub_clause_split: false.")
    elif "parent_section" in entry:
        failures.append(f"[G1] '{non_split_sample}' must not carry parent_section (non-split anchor).")

    return failures


def check_heading_hash_stability(mod):
    """Gate 2: same input -> same heading_hash per anchor across two builds."""
    failures = []
    eiaa_section_config = mod.load_section_config("eiaa")
    anchors_a = mod.build_anchors_from_config(config=eiaa_section_config["sections"])
    anchors_b = mod.build_anchors_from_config(config=eiaa_section_config["sections"])

    if anchors_a.keys() != anchors_b.keys():
        failures.append("[G2] Two builds from the same SECTION_CONFIG produced different anchor sets.")
        return failures

    for anchor in anchors_a:
        hash_a = anchors_a[anchor]["heading_hash"]
        hash_b = anchors_b[anchor]["heading_hash"]
        if hash_a != hash_b:
            failures.append(
                f"[G2] heading_hash for '{anchor}' is NOT stable across builds: {hash_a} != {hash_b}."
            )

    return failures


def check_docx_mode():
    """Gate 3: real-.docx mode resolves every SECTION_CONFIG anchor to a heading
    actually present in a synthetic fixture .docx (not a silent config fallback).

    python-docx is optional repo-wide (every other CI gate is stdlib-only; see
    build_anchor_map.py's own ImportError handling for --docx mode). If it is
    not installed, this gate is SKIPPED, not failed, so it never adds a new
    hard CI dependency. Returns (failures, skipped_message_or_None).
    """
    failures = []
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return failures, (
            "[G3] SKIPPED: python-docx not installed locally -- docx-mode fixture "
            "test not exercised (optional dependency, consistent with every other "
            "stdlib-only gate in this repo). Run `pip install python-docx` to "
            "exercise this gate."
        )

    import build_anchor_map as mod  # re-import for clarity within this scope

    # Explicitly "eiaa" (issue #343 repointed the registry default to the
    # small public "sample-agreement" sample playbook) -- see
    # check_section_config_fixture's docstring above.
    eiaa_sections = mod.load_section_config("eiaa")["sections"]

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "synthetic-standard-form.docx"
        doc = Document()
        for _anchor, heading, _sub_split, _parent in eiaa_sections:
            doc.add_heading(heading, level=2)
            doc.add_paragraph(f"Body text for {heading}.")
        doc.save(str(docx_path))

        anchors = mod.build_anchors_from_docx(docx_path, config=eiaa_sections)

        for anchor, heading, _sub_split, _parent in eiaa_sections:
            entry = anchors.get(anchor)
            if entry is None:
                failures.append(f"[G3] MISSING anchor '{anchor}' from docx-mode build.")
                continue
            if entry["heading"] != heading:
                failures.append(
                    f"[G3] '{anchor}' resolved heading {entry['heading']!r} does not match "
                    f"the fixture .docx heading {heading!r} -- docx-mode must read the real "
                    f"heading text from the document, not silently fall back to config."
                )

    return failures, None


def check_deterministic_artifact():
    """Gate 4: running the builder CLI twice on the same input produces a
    BYTE-IDENTICAL output file (not just a matching anchor_map_hash)."""
    failures = []
    builder_path = SCRIPTS_DIR / "build_anchor_map.py"

    with tempfile.TemporaryDirectory() as tmpdir:
        out_a = Path(tmpdir) / "run-a.anchor-map.json"
        out_b = Path(tmpdir) / "run-b.anchor-map.json"

        for out_path in (out_a, out_b):
            result = subprocess.run(
                [sys.executable, str(builder_path), "--output", str(out_path)],
                cwd=str(REPO_ROOT),
                capture_output=True,
                text=True,
            )
            if result.returncode != 0:
                failures.append(
                    f"[G4] builder invocation failed (exit {result.returncode}):\n{result.stderr}"
                )
                return failures

        bytes_a = out_a.read_bytes()
        bytes_b = out_b.read_bytes()

        if bytes_a != bytes_b:
            data_a = json.loads(bytes_a)
            data_b = json.loads(bytes_b)
            diffs = [
                key for key in data_a
                if key in data_b and data_a[key] != data_b[key]
            ]
            failures.append(
                "[G4] NOT DETERMINISTIC: two builder runs on identical input produced "
                "byte-different output files.\n"
                f"  Differing top-level keys: {diffs}\n"
                "  Issue #75 Green criterion: 'deterministic output (same input -> "
                "byte-identical map)'.\n"
                "  FIX: remove (or fix to the input, not wall-clock) any field that "
                "varies run-to-run -- e.g. a `generated_at` timestamp captured with "
                "datetime.now() -- from the written artifact, following the same "
                "no-timestamp-in-hashed-content convention as "
                "scripts/diff_standard_form.py's serialize_diff()/diff_hash()."
            )

    return failures


def main():
    mod, err = _import_builder()
    if mod is None:
        print(f"FAIL: {err}")
        sys.exit(1)

    all_failures = []
    all_failures += check_section_config_fixture(mod)
    all_failures += check_heading_hash_stability(mod)

    docx_failures, docx_skip_msg = check_docx_mode()
    all_failures += docx_failures

    all_failures += check_deterministic_artifact()

    if all_failures:
        print("FAIL: anchor-map builder fixture tests found the following issues:\n")
        for f in all_failures:
            print(f)
            print()
        print(f"Total failures: {len(all_failures)}")
        sys.exit(1)

    if docx_skip_msg:
        print(docx_skip_msg)

    print(
        "PASS: anchor-map builder fixture tests "
        "(section-config fixture, heading-hash stability, docx-mode resolution, "
        "byte-identical determinism)."
    )
    sys.exit(0)


if __name__ == "__main__":
    main()
