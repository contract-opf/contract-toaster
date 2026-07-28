#!/usr/bin/env python3
"""
One-shot generator for the issue #381 "reformatted contract" fixture.

NOT part of the test suite (no test_*.py name, not collected by check.sh) --
this is a build-time tool, run once to produce the committed
`reformatted-contract.SYNTHETIC.docx` under this directory, the same
"generate once, commit the output" convention
`tests/fixtures/gold_docx_204/_generate.py` uses for its own fixture pair.

Re-run this script (`python3 tests/fixtures/quote_redline_e2e/_generate.py`)
only if the fixture's content needs to change -- otherwise the committed
`.docx` IS the fixture.

## Why this fixture is NOT built from the standard form

Every other hand-built fixture in this repo that needs a "realistic" draft
(`tests/test_review_spine.py::_build_draft_docx`,
`tests/fixtures/gold_docx_204/_generate.py::build_draft_body_xml`) starts
from `diff_standard_form.load_standard_form_paragraphs()` and carries every
anchor's heading/text over VERBATIM except the one or two anchors it plants
a change at -- i.e. the draft still matches the standard form
section-for-section. Issue #381's whole point is the OPPOSITE: prove the
LLM-native quote-based path (issues #379/#380/#398) on a contract a
counterparty reformatted from scratch -- different article numbering
("ARTICLE IV" vs. the standard form's "sec-8"), different headings, merged/
reordered clauses, and paraphrased prose throughout, so that a structural
diff against the standard form could never anchor most of it. Before #380
retired the deterministic detector engine and standard-form diff from
issue-generation, a document shaped like this would have diffed almost
entirely as "no matching anchor" and the review would have fail-closed to
MANUAL_REVIEW_REQUIRED with no document produced at all (this ticket's own
Goal text). This fixture is therefore an ORIGINAL six-article contract, not
a copy of any `standard-forms/*.docx` paragraph.

Two deliberate deviations from Your Organization, LLC's standard playbook
position are planted, for the two issues the test's FakeBedrockClient primary
response reports:

  1. ARTICLE IV (Risk Allocation and Indemnification) -- unlimited one-way
     liability/indemnification in place of the standard mutual $150,000 cap
     (playbook topic `limitation-of-liability`). The test's model issue
     names an exact, unique, verbatim `source_quote` from this article --
     issue #379's quote-based patcher locates and applies it.
  2. ARTICLE II (Duration and Renewal) -- the non-renewal notice period is
     shortened from the standard 60 days to 10 days (playbook topic
     `term-length`). The test's model issue for THIS deviation carries a
     `source_quote` that PARAPHRASES the actual clause (a realistic
     imperfect-copy failure mode) rather than quoting it verbatim, so it
     legitimately does not locate -- proving the not-locatable ->
     flag-only path (`analysis_report.changes_not_applied`), never a
     silent/incorrect edit.

Uses python-docx (test-only dependency, matches
`tests/test_redline_quote_apply.py`'s convention) to build a real,
Word-authored-shaped `.docx` -- python-docx's default `Heading1` style
writes the exact `<w:pStyle w:val="Heading1"/>` this repo's other hand-built
fixtures use, so headings extract identically either way (confirmed against
`scripts/extraction_normalization_stage.py`'s heading-boundary detection).
"""

from __future__ import annotations

import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[3]
FIXTURE_DIR = Path(__file__).resolve().parent

sys.path.insert(0, str(REPO_ROOT))

import docx  # noqa: E402  (test-only dependency, requirements-dev.txt)

# ---------------------------------------------------------------------------
# Contract content. Every ARTICLE is a single python-docx paragraph (one
# `<w:p>`, one run) so each logical paragraph's extracted `text` is exactly
# the string below, character for character -- no run-splitting or sibling-
# joining surprises for the test's source_quote spans to account for.
# ---------------------------------------------------------------------------

TITLE = "CLINICAL EDUCATION PARTNERSHIP AGREEMENT"

PREAMBLE = (
    'This Clinical Placement and Affiliation Agreement (the "Agreement") is '
    'entered into by and between Riverside College of Health Sciences (the '
    '"Institution") and Your Organization, LLC (the "Company") to establish '
    "the terms on which Institution students will complete supervised "
    "clinical rotations at Company-affiliated training sites."
)

ARTICLE_I_HEADING = "ARTICLE I -- PROGRAM STRUCTURE AND OVERSIGHT"
ARTICLE_I_TEXT = (
    "The Company retains full and final authority over the acceptance, "
    "placement, and dismissal of any student referred by the Institution, "
    "exercised according to the Company's own reasonable, uniformly applied "
    "criteria. The Institution may, upon reasonable advance request, visit "
    "and inspect any training site at which its students are placed."
)

ARTICLE_II_HEADING = "ARTICLE II -- DURATION AND RENEWAL"
# NOTE: the standard position (playbook topic `term-length`) is 60 days'
# notice; this counterparty draft shortens it to 10 -- the planted deviation
# issue #2 (the test's NOT-LOCATABLE issue) targets, via a paraphrased quote
# that does not match this sentence verbatim.
ARTICLE_II_TEXT = (
    "This Agreement commences on the Effective Date and continues for an "
    "initial period of five (5) years. Thereafter, it shall renew "
    "automatically for successive one-year terms unless either party "
    "delivers written notice of non-renewal at least ten (10) days before "
    "the end of the then-current term. Any clinical placement already in "
    "progress at the time of a non-renewal shall be allowed to reach its "
    "natural completion notwithstanding the expiration of this Agreement."
)

ARTICLE_III_HEADING = "ARTICLE III -- STUDENT STATUS"
ARTICLE_III_TEXT = (
    "Students participating in the Program are not employees of the Company "
    "for any purpose and are not entitled to wages, benefits, or workers' "
    "compensation coverage through the Company; students remain covered, if "
    "at all, through the Institution's own arrangements. No payment shall be "
    "made by either party to the other, and no student shall receive "
    "remuneration from the Company on account of the Program."
)

ARTICLE_IV_HEADING = "ARTICLE IV -- RISK ALLOCATION AND INDEMNIFICATION"
# NOTE: the standard position (playbook topic `limitation-of-liability`) is a
# mutual $150,000 aggregate cap with a consequential-damages exclusion and no
# one-way indemnity. This counterparty draft removes the cap entirely and
# adds one-way indemnification running FROM the Company TO the Institution --
# the planted deviation issue #1 (the test's LOCATABLE issue) targets.
ARTICLE_IV_TEXT = (
    "Notwithstanding anything to the contrary elsewhere in this Agreement, "
    "Your Organization, LLC's liability under this Agreement shall be "
    "unlimited, and Your Organization, LLC shall defend, indemnify, and "
    "hold harmless the Institution, its trustees, officers, employees, and "
    "agents from and against any and all third-party claims, losses, "
    "damages, and expenses, including reasonable attorneys' fees, arising "
    "out of or in any way related to the Program, regardless of the degree "
    "of fault of any party."
)

ARTICLE_V_HEADING = "ARTICLE V -- CONFIDENTIALITY"
ARTICLE_V_TEXT = (
    "Each party shall use reasonable care to protect the other party's "
    "confidential information from unauthorized disclosure and shall use "
    "such information solely to carry out its obligations under this "
    "Agreement. This obligation does not extend to information that is or "
    "becomes publicly available through no fault of the receiving party, "
    "was already known to the receiving party without an obligation of "
    "confidence, or is independently developed."
)

ARTICLE_VI_HEADING = "ARTICLE VI -- GENERAL PROVISIONS"
ARTICLE_VI_TEXT = (
    "This Agreement, together with any written amendment signed by both "
    "parties, constitutes the entire agreement between the parties and "
    "supersedes all prior discussions. Nothing in this Agreement shall be "
    "construed to prevent either party from entering into similar "
    "arrangements with other institutions or organizations."
)


def build_document() -> "docx.Document":
    document = docx.Document()
    document.add_heading(TITLE, level=1)
    document.add_paragraph(PREAMBLE)

    for heading, text in (
        (ARTICLE_I_HEADING, ARTICLE_I_TEXT),
        (ARTICLE_II_HEADING, ARTICLE_II_TEXT),
        (ARTICLE_III_HEADING, ARTICLE_III_TEXT),
        (ARTICLE_IV_HEADING, ARTICLE_IV_TEXT),
        (ARTICLE_V_HEADING, ARTICLE_V_TEXT),
        (ARTICLE_VI_HEADING, ARTICLE_VI_TEXT),
    ):
        document.add_heading(heading, level=1)
        document.add_paragraph(text)

    return document


def main() -> None:
    document = build_document()
    docx_path = FIXTURE_DIR / "reformatted-contract.SYNTHETIC.docx"
    document.save(str(docx_path))
    print(f"Wrote {docx_path} ({docx_path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
