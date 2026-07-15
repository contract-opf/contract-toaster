#!/usr/bin/env python3
"""
Slice test (TDD) — real .docx standard-form loader + §10 sub-clause split
(issue #200).

## Root problem this proves fixed

standard-forms/README.md admitted the canonical eiaa-v1.0.0.docx was "to be
committed when the real .docx is ready". Until then
load_standard_form_paragraphs() synthesized the standard body from playbook
prose (diff_standard_form.py's synthetic mode) — never exercising real-docx
mode at all (zero .docx in the repo, python-docx optional) — and, worse,
real-docx mode COULD NOT have worked even if a real .docx existed: the
anchor map hand-splits §10 into four anchors with invented headings
("Miscellaneous: Notices", etc.), but the pre-#200 loader matched paragraphs
to anchors ONLY by heading text — and a real document has a single
"10. Miscellaneous" heading, so all four §10 anchors would resolve to empty
placeholder paragraphs, and every real draft would produce four phantom
"deleted" §10 hunks.

Per the 2026-07-10 owner decision (issue #200), the real EIAA form text is a
GC deliverable engineering cannot fabricate. This test instead drives the
REAL loader / anchor-map-builder / diff code over a clearly-labeled
SYNTHETIC placeholder — standard-forms/eiaa-v1.0.0.SYNTHETIC.docx (see
scripts/generate_synthetic_standard_form.py) — that faithfully expresses the
required *structure*, including the §10 sub-clause split under a single
real "Miscellaneous" heading (the exact case the pre-#200 heading-only
matcher failed). When the real .docx is provided, it drops in with NO code
change; only the fixture file changes.

## What this test asserts

  1. The loader parses real OOXML paragraphs from the .docx — not
     synthesized from playbook `exos_standard` prose.
  2. The anchor-map builder (real-docx mode) resolves the §10 sub-clause
     anchors from the real document structure: no empty placeholder
     paragraphs, no phantom "deleted" §10 hunks.
  3. A self-diff of the placeholder against itself is all-unchanged, with
     zero deleted/placeholder hunks anywhere in the form.
  4. Heading/anchor resolution works against the real numbering style: a
     plain heading with no manually-typed section number (as in the
     SYNTHETIC fixture) resolves correctly, while a manually-typed number
     in the heading text (e.g. "1.2 Admitting Students") — which a hand-typed
     real .docx might use instead of genuine Word auto-numbering — does NOT
     match and is a visible drift signal (empty placeholder paragraph), not
     a silent success.

python-docx is a REQUIRED dependency for this test (declared in
requirements-dev.txt) — unlike other docx-mode gates in this repo, this
test's whole point is exercising real-docx mode, so it must not silently
skip.

Exit codes: 0 = pass, 1 = fail
"""

import sys
import tempfile
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
SCRIPTS_DIR = REPO_ROOT / "scripts"
SYNTHETIC_DOCX = REPO_ROOT / "standard-forms" / "eiaa-v1.0.0.SYNTHETIC.docx"

sys.path.insert(0, str(SCRIPTS_DIR))


def _import_modules():
    try:
        import diff_standard_form as dsf  # type: ignore
        import build_anchor_map as bam  # type: ignore
        return dsf, bam, ""
    except ImportError as exc:
        return None, None, (
            f"MISSING: scripts/diff_standard_form.py or "
            f"scripts/build_anchor_map.py does not import ({exc})."
        )


def main():
    failures = []

    dsf, bam, err = _import_modules()
    if dsf is None:
        print(f"FAIL: {err}")
        sys.exit(1)

    try:
        from docx import Document  # type: ignore
    except ImportError:
        print(
            "FAIL: python-docx is not installed. It is a REQUIRED dev "
            "dependency for this test (see requirements-dev.txt) -- unlike "
            "other docx-mode gates in this repo, this test's whole point is "
            "exercising real-docx mode against the SYNTHETIC placeholder "
            "standard form, so it must not silently skip.\n"
            "  FIX: pip install -r requirements-dev.txt"
        )
        sys.exit(1)

    if not SYNTHETIC_DOCX.exists():
        print(
            f"FAIL: SYNTHETIC placeholder standard form not found at "
            f"{SYNTHETIC_DOCX}.\n"
            f"  FIX: python3 scripts/generate_synthetic_standard_form.py"
        )
        sys.exit(1)

    # -------------------------------------------------------------------
    # Setup: load the real-docx-mode standard-form paragraphs, and the
    # synthetic (playbook-prose) standard-form paragraphs, for comparison.
    # -------------------------------------------------------------------
    docx_standard = dsf.load_standard_form_paragraphs(docx_path=SYNTHETIC_DOCX)
    synthetic_standard = dsf.load_standard_form_paragraphs()

    docx_by_anchor = {p["anchor"]: p for p in docx_standard}
    synthetic_by_anchor = {p["anchor"]: p for p in synthetic_standard}

    # --- G1: real OOXML paragraphs, not synthesized from playbook prose ----
    # Every anchor covered by a playbook topic has DIFFERENT text in docx
    # mode than in synthetic mode (docx mode reads the .docx's own prose;
    # synthetic mode reads playbooks/eiaa-v1.0.0.json's exos_standard
    # field) -- if they were ever equal for a topic-covered anchor, that
    # would mean docx mode silently fell back to playbook prose instead of
    # the real document.
    same_text_anchors = []
    for anchor, docx_para in docx_by_anchor.items():
        if docx_para.get("structural") or docx_para.get("absent_from_form"):
            continue
        synth_para = synthetic_by_anchor.get(anchor)
        if synth_para is None:
            continue
        if not docx_para["text"]:
            continue
        if docx_para["text"] == synth_para["text"]:
            same_text_anchors.append(anchor)

    if same_text_anchors:
        failures.append(
            "[G1] docx-mode text is IDENTICAL to synthetic (playbook-prose) "
            f"text for anchor(s) {same_text_anchors} -- the real-docx loader "
            "must read paragraph text from the .docx itself, never fall "
            "back to playbook prose."
        )

    # Spot-check: a known .docx-only phrase (from
    # scripts/generate_synthetic_standard_form.py's BODY_TEXT, never present
    # in the playbook's exos_standard prose) must appear.
    sec8_docx_text = docx_by_anchor.get("sec-8", {}).get("text", "")
    if "Neither party shall be liable to the other for consequential" not in sec8_docx_text:
        failures.append(
            f"[G1b] sec-8 docx-mode text does not contain the expected "
            f".docx-sourced phrase. Got: {sec8_docx_text!r}"
        )

    # --- G2: §10 sub-clause split resolved from real document structure ----
    sec10_anchors = [
        "sec-10-notices",
        "sec-10-non-exclusive",
        "sec-10-merger",
        "sec-10-precedence",
    ]
    expected_markers = {
        "sec-10-notices": "(a)",
        "sec-10-non-exclusive": "(b)",
        "sec-10-merger": "(c)",
        "sec-10-precedence": "(d)",
    }
    for anchor in sec10_anchors:
        para = docx_by_anchor.get(anchor)
        if para is None:
            failures.append(f"[G2] MISSING §10 sub-clause anchor '{anchor}' in docx-mode standard form.")
            continue
        if not para["text"].strip():
            failures.append(
                f"[G2] '{anchor}' has an EMPTY placeholder paragraph in docx mode -- "
                "the pre-#200 bug this test exists to catch (all four §10 anchors "
                "collapsing to empty placeholders because the real document has a "
                "single '10. Miscellaneous' heading, not four)."
            )
            continue
        marker = expected_markers[anchor]
        if not para["text"].startswith(marker):
            failures.append(
                f"[G2b] '{anchor}' text does not start with its expected marker "
                f"{marker!r}. Got: {para['text'][:60]!r}"
            )

    # Anchor-map builder (real-docx mode): every §10 sub-clause anchor must
    # resolve without falling back to a bogus per-anchor heading match.
    anchors_from_builder = bam.build_anchors_from_docx(SYNTHETIC_DOCX)
    for anchor in sec10_anchors:
        entry = anchors_from_builder.get(anchor)
        if entry is None:
            failures.append(f"[G2c] MISSING §10 sub-clause anchor '{anchor}' from anchor-map builder docx mode.")
            continue
        if entry.get("sub_clause_split") is not True or entry.get("parent_section") != "sec-10":
            failures.append(
                f"[G2d] '{anchor}' from anchor-map builder docx mode must have "
                f"sub_clause_split=True, parent_section='sec-10' (got {entry!r})."
            )

    # --- G3: self-diff is all-unchanged, zero deleted/placeholder hunks ----
    self_draft = [{"heading": p["heading"], "text": p["text"]} for p in docx_standard]
    self_hunks = dsf.diff_draft_against_standard(docx_standard, self_draft)

    non_unchanged = [h for h in self_hunks if h["kind"] != "unchanged"]
    if non_unchanged:
        failures.append(
            "[G3] Self-diff of the SYNTHETIC placeholder against itself is NOT "
            f"all-unchanged. Non-unchanged hunks: {non_unchanged}"
        )

    deleted_hunks = [h for h in self_hunks if h["kind"] == "deleted"]
    if deleted_hunks:
        failures.append(
            f"[G3b] Self-diff produced {len(deleted_hunks)} 'deleted' hunk(s) -- "
            "these are exactly the phantom hunks the pre-#200 empty-placeholder "
            f"bug produced for §10. Anchors: {[h['anchor'] for h in deleted_hunks]}"
        )

    sec10_hunks = [h for h in self_hunks if h["anchor"] in sec10_anchors]
    if len(sec10_hunks) != 4:
        failures.append(
            f"[G3c] Expected exactly 4 §10 sub-clause hunks in the self-diff "
            f"(one per sub-anchor), got {len(sec10_hunks)}: {sec10_hunks}"
        )
    for h in sec10_hunks:
        if h["kind"] != "unchanged":
            failures.append(
                f"[G3d] §10 sub-clause hunk for '{h['anchor']}' is kind={h['kind']!r}, "
                "expected 'unchanged' in a self-diff."
            )

    # --- G4: heading resolution works against real numbering style ---------
    # G4a: the SYNTHETIC fixture's plain (no manually-typed number) headings
    # resolve correctly -- already proven by G1/G2 above finding non-empty
    # text for ordinary anchors; spot-check one explicitly here.
    admitting_students = docx_by_anchor.get("sec-1.2")
    if admitting_students is None or not admitting_students["text"].strip():
        failures.append(
            "[G4a] 'sec-1.2' (heading 'Admitting Students', no manually-typed "
            "section number) did not resolve to a non-empty paragraph in docx mode."
        )

    # G4b: a manually-typed section number IN the heading text (as a
    # hand-typed real .docx might do, instead of genuine Word
    # auto-numbering) does NOT match the anchor map's plain heading string
    # -- this must be a visible drift signal (empty placeholder), not a
    # silent success, per issue #200's explicit warning about this failure
    # mode.
    with tempfile.TemporaryDirectory() as tmpdir:
        bad_docx_path = Path(tmpdir) / "manually-numbered.docx"
        doc = Document()
        doc.add_heading("1.2 Admitting Students", level=2)
        doc.add_paragraph("Students are admitted per your criteria.")
        doc.save(str(bad_docx_path))

        anchor_map_data = dsf._load_active_anchor_map()
        anchors = anchor_map_data["anchors"]
        sub_clause_splits = anchor_map_data.get("sub_clause_splits", {})
        bad_standard = dsf._load_standard_form_paragraphs_from_docx(
            bad_docx_path, anchors, sub_clause_splits
        )
        bad_by_anchor = {p["anchor"]: p for p in bad_standard}
        bad_sec12 = bad_by_anchor.get("sec-1.2")
        if bad_sec12 is None:
            failures.append("[G4b] 'sec-1.2' missing entirely from manually-numbered-heading fixture result.")
        elif bad_sec12["text"].strip():
            failures.append(
                "[G4b] A manually-typed section number in the heading text "
                "('1.2 Admitting Students') UNEXPECTEDLY matched the anchor map's "
                "plain heading ('Admitting Students') -- heading matching must be "
                "exact (whitespace/case-normalized only), so this drift mode stays "
                "visible instead of silently 'working'."
            )

    # --- Report --------------------------------------------------------------
    if failures:
        print("FAIL: standard-form loader / §10 sub-clause split gate (issue #200).\n")
        for f in failures:
            print(f)
            print()
        print(f"Total failures: {len(failures)}")
        sys.exit(1)
    else:
        print(
            "PASS: standard-form loader / §10 sub-clause split gate (issue #200). "
            f"{len(docx_standard)} docx-mode paragraphs, {len(self_hunks)} self-diff "
            "hunks, all unchanged, zero deleted/placeholder hunks."
        )
        sys.exit(0)


if __name__ == "__main__":
    main()
