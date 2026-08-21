#!/usr/bin/env python3
"""
tools/churn_docx.py -- issue #565: synthetic churn generator.

## Why this exists

`tools/document_spine_smoke.py` (this issue's other half) is a discovery
harness for the model-free document spine (extract -> normalize -> locate ->
apply). Pointed at a private corpus it finds real structural failure
classes; pointed at nothing, it proves nothing, because the public repo
cannot commit that corpus. This module is the other half of the fix: a
generator that manufactures a small, committable, negotiation-shaped
synthetic corpus reproducing every structural failure class the private
corpus has discovered so far, so the public regression gate
(`tests/test_document_shapes.py`) can prove the spine still survives each
one -- forever, offline, without ever reading a real agreement.

## No network, no vendored documents

Every base contract is generated PROGRAMMATICALLY with `python-docx`
(already a pinned dev dependency -- see requirements-dev.txt's own
docstring). No network downloads, no vendored third-party `.docx` files:
determinism and licensing both (this issue's Out of scope).

## Two layers: build, then churn

`build_base_document()` uses python-docx for what it is actually good at --
producing a real, well-formed base `.docx` with headings, numbered clauses,
and a table -- with NO attempt to express tracked changes, custom namespace
prefixes, or split/stripped paragraph structure through python-docx's object
model, because it has no API surface for any of that.

Every named transform in `TRANSFORMS` below instead operates directly on
the built document's `word/document.xml`, via the SAME zipfile+ElementTree,
root-namespace-preservation technique this repo's own redline modules use
(`redline_inplace._root_open_tag` / `register_declared_namespaces` /
`_merge_hoisted_namespaces`, reused here rather than reimplemented -- see
`_rewrite_document_xml` below) -- every other zip part (styles.xml,
docProps/*, ...) survives byte-for-byte; only `word/document.xml` is
rewritten.

## Determinism

Every transform is a pure function of `(docx_bytes, seed)`: the same input
bytes and the same seed always produce the same output bytes. `seed`
currently only selects among a small fixed set of synthetic author-name
pairs (`tracked_changes_multi_author`, `nested_ins_del`) -- it exists so a
caller can vary WHICH fabricated names appear without losing reproducibility,
not to introduce randomness.

## The seven transforms (one per issue #565 Scope, plus #530's follow-up)

  - `tracked_changes_multi_author` -- two different authors' `<w:ins>`/
    `<w:del>` clusters, back-to-back with no intervening plain text, on the
    Indemnification clause. Exercises issue #563's narrowed accept-all rule
    (more than one pending cluster/author no longer fails closed).
  - `curly_punctuation` -- every straight `'`/`"`/` - ` in every `<w:t>` run
    becomes its Word-autocorrect typographic equivalent. Exercises
    `quote_locate.py`'s `_TYPOGRAPHIC_FOLD` table (15 of 16 documents in the
    real EIAA corpus carry curly punctuation -- see that module's
    docstring).
  - `split_paragraphs` -- the Definitions clause's three sentences become
    three sibling `<w:p>` elements instead of one. Exercises issue #564's
    `physical_spans` paragraph-join accounting.
  - `strip_heading_styles` -- every `Heading*`-style `<w:pStyle>` is
    removed. Exercises `clause_boundaries.py`'s document-signals fallback
    (every heading here is ALSO a numbered lead-in, e.g. "1. Definitions",
    so the fallback still fires with the style gone).
  - `reserved_ns_prefix` -- declares `xmlns:ns0="..."` on the document
    root. The exact issue #560/#561 trigger (`ET.register_namespace`
    refuses any `ns<digits>` prefix); fixed on main, this shape exists so
    the fix can never silently regress.
  - `nested_ins_del` -- an insertion (`<w:ins>`) later itself deleted
    (`<w:del>` nested inside it) on the Confidentiality clause -- a
    net-zero edit a real negotiation history routinely contains.
  - `pending_change_inside_field_code` (issue #530 follow-up) -- a
    `<w:fldSimple>` cross-reference field on the Term and Termination
    clause whose cached-result region carries a pending `<w:del>`/`<w:ins>`
    tracked change -- the counterparty is live-editing which section number
    the field displays. Exercises issue #530's narrowed accept-all rule: a
    pending change `inside_field_code` no longer fails closed on its own;
    watched failing first against this REAL document shape (a hand-built
    `normalize_input.py` dict fixture cannot exercise the extractor's own
    field-code detection at all -- see `tests/redline/test_normalize_
    pending_tracked_changes.py`'s `pending_change_inside_field_code.json`
    for that unit-level half, which this generated shape complements rather
    than replaces).

## What this is NOT

No model calls, no review quality of any kind (this issue's Out of scope).
No LibreOffice round-trip transform -- `soffice` is not a repo dependency;
see `docs/document-spine-smoke.md` for that as a manual, human-run extra.

See also: `scripts/extraction_normalization_stage.py`,
`scripts/quote_locate.py`, `scripts/redline_quote_apply.py`,
`tests/fixtures/adversarial/` and `tests/test_reserved_namespace_prefix_
560.py` for the same generated-fixture / zipfile+ElementTree conventions
this module follows.
"""

from __future__ import annotations

import argparse
import io
import sys
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from typing import Callable

REPO_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = REPO_ROOT / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import redline_inplace as ri  # noqa: E402 -- root-namespace-preservation helpers, reused not reimplemented

try:
    from docx import Document  # python-docx -- requirements-dev.txt
except ImportError as _exc:  # pragma: no cover -- exercised only in a broken venv
    Document = None  # type: ignore[assignment]
    _DOCX_IMPORT_ERROR: Exception | None = _exc
else:
    _DOCX_IMPORT_ERROR = None

WORD_NS = ri.WORD_NS
_XML_NS = "http://www.w3.org/XML/1998/namespace"


def _w(tag: str) -> str:
    return f"{{{WORD_NS}}}{tag}"


# ---------------------------------------------------------------------------
# Base document content -- every string below is fabricated. No real party
# names, no real agreement text, no vendored third-party paper.
#
# Every base contract shares these SAME five clause bodies so a transform
# can find its target clause by stable, known text regardless of which base
# contract ("flavor") is being churned -- only the title, one extra clause,
# and the table vary per flavor.
# ---------------------------------------------------------------------------

DEFINITIONS_BODY = (
    "Confidential Information means any non-public information disclosed by "
    "either party under this Agreement. Recipient means the party receiving "
    "Confidential Information from the Discloser. Effective Date means the "
    "date this Agreement is signed by both parties."
)
CONFIDENTIALITY_BODY = (
    "Each party shall protect the other party's Confidential Information "
    "using reasonable efforts and shall not disclose it to any third party "
    "without prior written consent."
)
INDEMNIFICATION_BODY = "The Provider shall indemnify the Recipient against all claims whatsoever."
TERM_BODY = (
    "The Receiving Party's obligations under this Section 4 survive "
    "termination for a period of three (3) years - unless earlier "
    "terminated by written notice."
)
GOVERNING_LAW_BODY = (
    "This Agreement is governed by the laws of the State of Delaware, "
    "without regard to its conflict of laws principles."
)

# The quote every shape test locates and patches to prove the document
# still survives the full spine. Deliberately left untouched by every
# transform below (no apostrophe/quote/dash, never a transform's target),
# so it is safe to search for verbatim regardless of which transform
# produced the shape.
ANCHOR_QUOTE = GOVERNING_LAW_BODY

_STANDARD_CLAUSES: list[tuple[str, str]] = [
    ("1. Definitions", DEFINITIONS_BODY),
    ("2. Confidentiality", CONFIDENTIALITY_BODY),
    ("3. Indemnification", INDEMNIFICATION_BODY),
    ("4. Term and Termination", TERM_BODY),
    ("5. Governing Law", GOVERNING_LAW_BODY),
]

BASE_DOCUMENT_SPECS: dict[str, dict] = {
    "mutual-nda": {
        "title": "MUTUAL NON-DISCLOSURE AGREEMENT",
        "extra_clause": (
            "6. Return of Materials",
            "Upon request, each party shall return or destroy all documents "
            "containing the other party's Confidential Information.",
        ),
        "table": [
            ["Party", "Notice Address"],
            ["Provider", "500 Synthetic Ave, Suite 100"],
            ["Recipient", "800 Fabricated Blvd, Suite 200"],
        ],
    },
    "master-services-agreement": {
        "title": "MASTER SERVICES AGREEMENT",
        "extra_clause": (
            "6. Fees and Payment",
            "Customer shall pay all undisputed fees within thirty (30) days "
            "of receipt of an accurate invoice.",
        ),
        "table": [
            ["Service", "Monthly Fee"],
            ["Implementation", "$5,000"],
            ["Support", "$1,500"],
        ],
    },
    "data-processing-addendum": {
        "title": "DATA PROCESSING ADDENDUM",
        "extra_clause": (
            "6. Sub-processors",
            "Processor may engage sub-processors provided it maintains a "
            "current list and imposes materially equivalent obligations.",
        ),
        "table": [
            ["Sub-processor", "Processing Location"],
            ["Synthetic Hosting Co.", "Ireland"],
            ["Fabricated Analytics Inc.", "United States"],
        ],
    },
    "statement-of-work": {
        "title": "STATEMENT OF WORK",
        "extra_clause": (
            "6. Deliverables",
            "Vendor shall deliver each milestone described in Exhibit A on "
            "the schedule set out therein.",
        ),
        "table": [
            ["Milestone", "Due Date"],
            ["Design Review", "2026-09-01"],
            ["Final Delivery", "2026-11-01"],
        ],
    },
}


def build_base_document(spec_name: str) -> bytes:
    """Builds one small synthetic contract (a title, five standard clauses
    plus one flavor-specific clause -- every heading a numbered lead-in --
    plus a table) with python-docx. Fully synthetic and fabricated; see
    module docstring.
    """
    if Document is None:  # pragma: no cover
        raise RuntimeError(f"python-docx is required to build base documents: {_DOCX_IMPORT_ERROR}")
    if spec_name not in BASE_DOCUMENT_SPECS:
        raise ValueError(f"unknown base document spec: {spec_name!r} (known: {sorted(BASE_DOCUMENT_SPECS)})")
    spec = BASE_DOCUMENT_SPECS[spec_name]

    doc = Document()
    doc.add_paragraph(spec["title"])
    for heading, body in _STANDARD_CLAUSES:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    extra_heading, extra_body = spec["extra_clause"]
    doc.add_heading(extra_heading, level=1)
    doc.add_paragraph(extra_body)

    rows = spec["table"]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            table.cell(r, c).text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Shared OOXML surgery helpers -- the zipfile+ElementTree layer every
# transform below is built on.
# ---------------------------------------------------------------------------


def _rewrite_document_xml(docx_bytes: bytes, mutate: Callable[[ET.Element], None]) -> bytes:
    """Reparses `word/document.xml`, calls `mutate(root)` (which mutates the
    tree in place), and reserializes -- every other zip part survives
    byte-for-byte, and every original root `xmlns` declaration survives too.

    Same root-namespace-preservation technique `redline_quote_apply.
    _rewrite_revision_dates` / `redline_generate.
    inject_export_marker_and_footnotes` use (issue #560/#561): a straight
    `ET.fromstring` -> mutate -> `ET.tostring` round trip silently drops any
    root `xmlns` declaration ElementTree does not see used, which is exactly
    what a python-docx base document's document.xml is full of (a dozen
    declared-but-unused drawing/VML namespaces). Reused via
    `redline_inplace`, not reimplemented.
    """
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        infos = zf.infolist()
        originals = {info.filename: zf.read(info.filename) for info in infos}

    doc_xml_text = originals[ri.DOCUMENT_PART].decode("utf-8")
    original_root_open_tag = ri._root_open_tag(doc_xml_text)
    ri.register_declared_namespaces(ri._declared_namespaces_anywhere(doc_xml_text))

    root = ET.fromstring(originals[ri.DOCUMENT_PART])
    mutate(root)

    serialized = ET.tostring(root, encoding="unicode")
    auto_root_open_tag = ri._root_open_tag(serialized)
    body_and_close = serialized[len(auto_root_open_tag) :]
    root_open_tag = ri._merge_hoisted_namespaces(original_root_open_tag, auto_root_open_tag)
    new_document_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        + root_open_tag.encode("utf-8")
        + body_and_close.encode("utf-8")
    )

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf_out:
        for info in infos:
            data = new_document_xml if info.filename == ri.DOCUMENT_PART else originals[info.filename]
            zf_out.writestr(info, data)
    return out.getvalue()


def _find_paragraph_by_text(root: ET.Element, needle: str) -> ET.Element:
    """The (first) `<w:p>` under `<w:body>` whose joined `<w:t>` text
    contains `needle` -- every base document's clause bodies are unique
    strings, so "contains" is sufficient and avoids requiring exact
    whitespace-for-whitespace equality with python-docx's own output."""
    body = root.find(_w("body"))
    if body is not None:
        for p_el in body.iter(_w("p")):
            text = "".join(t.text or "" for t in p_el.iter(_w("t")))
            if needle in text:
                return p_el
    raise ValueError(f"no paragraph found containing the expected clause text (length {len(needle)})")


def _parent_of(root: ET.Element, target: ET.Element) -> ET.Element:
    for parent in root.iter():
        if target in list(parent):
            return parent
    raise ValueError("element has no parent in this tree")


def _clear_runs(p_el: ET.Element) -> None:
    for child in list(p_el):
        if child.tag == _w("r"):
            p_el.remove(child)


def _append_run(parent: ET.Element, text: str) -> ET.Element:
    """Appends a `<w:r><w:t xml:space="preserve">text</w:t></w:r>` to
    `parent` (a `<w:p>`, `<w:ins>`, or `<w:del>` element -- all three may
    directly contain a run)."""
    r = ET.SubElement(parent, _w("r"))
    t = ET.SubElement(r, _w("t"))
    t.set(f"{{{_XML_NS}}}space", "preserve")
    t.text = text
    return r


_SYNTHETIC_AUTHOR_PAIRS: list[tuple[str, str]] = [
    ("Alice Counsel", "Bob Counterparty"),
    ("Priya Legal", "Marcus Ops"),
]


def _author_pair(seed: int) -> tuple[str, str]:
    return _SYNTHETIC_AUTHOR_PAIRS[seed % len(_SYNTHETIC_AUTHOR_PAIRS)]


# ---------------------------------------------------------------------------
# Named transforms -- each independently toggleable, each a pure function
# of (docx_bytes, seed).
# ---------------------------------------------------------------------------


def tracked_changes_multi_author(docx_bytes: bytes, *, seed: int = 0) -> bytes:
    """Rewrites the Indemnification clause as: plain lead-in, then a
    `<w:del>` from one author immediately followed by a `<w:ins>` from a
    SECOND author (no intervening plain run) -- two distinct clusters by
    `extraction_normalization_stage._ParaBuilder._ensure_cluster`'s own
    rule, from two different authors. Exercises issue #563's narrowed
    accept-all rule: more than one pending cluster/author no longer fails
    closed on its own.
    """
    author_a, author_b = _author_pair(seed)

    def mutate(root: ET.Element) -> None:
        p_el = _find_paragraph_by_text(root, INDEMNIFICATION_BODY)
        _clear_runs(p_el)
        _append_run(p_el, "The Provider shall indemnify the Recipient against ")

        del_el = ET.SubElement(p_el, _w("del"))
        del_el.set(_w("id"), str(1000 + seed))
        del_el.set(_w("author"), author_a)
        del_el.set(_w("date"), "2026-08-01T00:00:00Z")
        del_run = ET.SubElement(del_el, _w("r"))
        del_text = ET.SubElement(del_run, _w("delText"))
        del_text.text = "all claims whatsoever"

        ins_el = ET.SubElement(p_el, _w("ins"))
        ins_el.set(_w("id"), str(1001 + seed))
        ins_el.set(_w("author"), author_b)
        ins_el.set(_w("date"), "2026-08-02T00:00:00Z")
        _append_run(ins_el, "third-party claims arising from any breach of this Agreement")

        _append_run(p_el, ".")

    return _rewrite_document_xml(docx_bytes, mutate)


# Word autocorrects a typed apostrophe/quote/dash to its typographic
# equivalent -- see `scripts/quote_locate.py`'s `_TYPOGRAPHIC_FOLD` table
# docstring: 15 of 16 normalizable documents in the real EIAA corpus carry
# curly punctuation. This transform reproduces exactly that encoding-only
# divergence, nothing else.
def curly_punctuation(docx_bytes: bytes, *, seed: int = 0) -> bytes:
    """Converts every straight `'`, `"`, and space-flanked `-` in every
    `<w:t>` run to Word's typographic auto-correct equivalent (`'`, `"`
    alternating open/close, en dash)."""
    state = {"double_open": True}

    def convert(text: str) -> str:
        out_chars: list[str] = []
        for ch in text:
            if ch == "'":
                out_chars.append("’")
            elif ch == '"':
                out_chars.append("“" if state["double_open"] else "”")
                state["double_open"] = not state["double_open"]
            else:
                out_chars.append(ch)
        return "".join(out_chars).replace(" - ", " – ")

    def mutate(root: ET.Element) -> None:
        for t_el in root.iter(_w("t")):
            if t_el.text:
                t_el.text = convert(t_el.text)

    return _rewrite_document_xml(docx_bytes, mutate)


def split_paragraphs(docx_bytes: bytes, *, seed: int = 0) -> bytes:
    """Splits the Definitions clause's three sentences across three sibling
    `<w:p>` elements (replacing the original single paragraph in place) --
    exercises issue #564's `physical_spans` accounting: a quote confined to
    one sibling still locates as `found`; a quote spanning two siblings
    locates as `spans_paragraph_break`, not a false `not_found`.
    """

    def mutate(root: ET.Element) -> None:
        p_el = _find_paragraph_by_text(root, DEFINITIONS_BODY)
        parent = _parent_of(root, p_el)
        index = list(parent).index(p_el)
        parent.remove(p_el)

        sentences = [s.strip() for s in DEFINITIONS_BODY.split(". ") if s.strip()]
        sentences = [s if s.endswith(".") else f"{s}." for s in sentences]
        for offset, sentence in enumerate(sentences):
            new_p = ET.Element(_w("p"))
            _append_run(new_p, sentence)
            parent.insert(index + offset, new_p)

    return _rewrite_document_xml(docx_bytes, mutate)


def strip_heading_styles(docx_bytes: bytes, *, seed: int = 0) -> bytes:
    """Removes every `Heading*`-style `<w:pStyle>` document-wide, so
    `clause_boundaries.py`'s document-signals fallback (numbered lead-in,
    since every heading here is literally "N. Title") must do the boundary
    detection that the (now absent) style would otherwise have done.
    """

    def mutate(root: ET.Element) -> None:
        body = root.find(_w("body"))
        if body is None:
            return
        for p_el in body.iter(_w("p")):
            ppr = p_el.find(_w("pPr"))
            if ppr is None:
                continue
            pstyle = ppr.find(_w("pStyle"))
            if pstyle is None:
                continue
            val = (pstyle.get(_w("val")) or "").strip().lower()
            if val.startswith("heading"):
                ppr.remove(pstyle)

    return _rewrite_document_xml(docx_bytes, mutate)


def reserved_ns_prefix(docx_bytes: bytes, *, seed: int = 0) -> bytes:
    """Declares `xmlns:ns0="..."` on the document root -- the exact
    structural trigger for issue #560/#561 (`ET.register_namespace` refuses
    any `ns<digits>` prefix; one call site in the redline path did not guard
    that, so 65% of a real 31-document corpus crashed before ever reaching
    this stage's normal reason codes). Fixed on main; this shape exists so
    the fix can never silently regress.

    A pure string splice on the root open tag, deliberately NOT routed
    through `_rewrite_document_xml`'s ElementTree round trip: ElementTree's
    own `register_namespace` raises `ValueError` on exactly this pattern
    (see `redline_inplace.register_declared_namespaces`'s docstring), so
    writing the declaration as literal text is the only way to produce this
    shape at all.
    """
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        infos = zf.infolist()
        originals = {info.filename: zf.read(info.filename) for info in infos}

    doc_xml_text = originals[ri.DOCUMENT_PART].decode("utf-8")
    open_tag = ri._root_open_tag(doc_xml_text)
    new_open_tag = open_tag[:-1] + ' xmlns:ns0="http://schemas.example.com/churn">'
    new_doc_xml_text = doc_xml_text.replace(open_tag, new_open_tag, 1)
    originals[ri.DOCUMENT_PART] = new_doc_xml_text.encode("utf-8")

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf_out:
        for info in infos:
            zf_out.writestr(info, originals[info.filename])
    return out.getvalue()


def nested_ins_del(docx_bytes: bytes, *, seed: int = 0) -> bytes:
    """Rewrites the Confidentiality clause so one author's insertion
    (`<w:ins>`) is itself later deleted by a second author (`<w:del>`
    nested inside it) -- a net-zero edit ("commercially " typed, then
    struck) that real negotiation history routinely contains. The deleted
    text never reaches `resulting_text` regardless of nesting depth, so
    this clause's accepted text should come back byte-for-byte equal to
    `CONFIDENTIALITY_BODY`.
    """
    author_a, author_b = _author_pair(seed)

    def mutate(root: ET.Element) -> None:
        p_el = _find_paragraph_by_text(root, CONFIDENTIALITY_BODY)
        _clear_runs(p_el)
        _append_run(p_el, "Each party shall protect the other party's Confidential Information using ")

        ins_el = ET.SubElement(p_el, _w("ins"))
        ins_el.set(_w("id"), str(2000 + seed))
        ins_el.set(_w("author"), author_a)
        ins_el.set(_w("date"), "2026-08-01T00:00:00Z")
        del_el = ET.SubElement(ins_el, _w("del"))
        del_el.set(_w("id"), str(2001 + seed))
        del_el.set(_w("author"), author_b)
        del_el.set(_w("date"), "2026-08-02T00:00:00Z")
        del_run = ET.SubElement(del_el, _w("r"))
        del_text = ET.SubElement(del_run, _w("delText"))
        del_text.text = "commercially "

        _append_run(
            p_el,
            "reasonable efforts and shall not disclose it to any third party without prior written consent.",
        )

    return _rewrite_document_xml(docx_bytes, mutate)


def pending_change_inside_field_code(docx_bytes: bytes, *, seed: int = 0) -> bytes:
    """Rewrites the Term and Termination clause so a `<w:fldSimple>`
    cross-reference field's cached-result region carries a PENDING tracked
    change -- the counterparty is live-editing which section number the
    field displays (a `<w:del>` of the old section reference immediately
    followed by an `<w:ins>` of the new one, both from the same author,
    inside the field). Exercises issue #530's narrowed accept-all rule: a
    pending change whose `inside_field_code` is true no longer fails closed
    on its own -- `extraction_normalization_stage._process_fld_simple`
    bubbles it up as an ordinary cluster, and that cluster's own
    `resulting_text` is already the paragraph's resolved text with the
    field's edit folded in, same as any other pending revision.
    """
    author, _ = _author_pair(seed)

    def mutate(root: ET.Element) -> None:
        p_el = _find_paragraph_by_text(root, TERM_BODY)
        _clear_runs(p_el)
        _append_run(p_el, "The Receiving Party's obligations under ")

        fld_el = ET.SubElement(p_el, _w("fldSimple"))
        fld_el.set(_w("instr"), " REF SectionFour \\h ")

        del_el = ET.SubElement(fld_el, _w("del"))
        del_el.set(_w("id"), str(4000 + seed))
        del_el.set(_w("author"), author)
        del_el.set(_w("date"), "2026-08-01T00:00:00Z")
        del_run = ET.SubElement(del_el, _w("r"))
        del_text = ET.SubElement(del_run, _w("delText"))
        del_text.text = "this Section 4"

        ins_el = ET.SubElement(fld_el, _w("ins"))
        ins_el.set(_w("id"), str(4001 + seed))
        ins_el.set(_w("author"), author)
        ins_el.set(_w("date"), "2026-08-02T00:00:00Z")
        _append_run(ins_el, "this Section 5")

        _append_run(
            p_el,
            " survive termination for a period of three (3) years - unless "
            "earlier terminated by written notice.",
        )

    return _rewrite_document_xml(docx_bytes, mutate)


TRANSFORMS: dict[str, Callable[..., bytes]] = {
    "tracked_changes_multi_author": tracked_changes_multi_author,
    "curly_punctuation": curly_punctuation,
    "split_paragraphs": split_paragraphs,
    "strip_heading_styles": strip_heading_styles,
    "reserved_ns_prefix": reserved_ns_prefix,
    "nested_ins_del": nested_ins_del,
    "pending_change_inside_field_code": pending_change_inside_field_code,
}


def apply_transform(docx_bytes: bytes, name: str, *, seed: int = 0) -> bytes:
    if name not in TRANSFORMS:
        raise ValueError(f"unknown transform: {name!r} (known: {sorted(TRANSFORMS)})")
    return TRANSFORMS[name](docx_bytes, seed=seed)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main(argv: list[str] | None = None) -> int:  # pragma: no cover - manual/CLI entry point
    parser = argparse.ArgumentParser(
        description=(
            "Manufacture a synthetic, negotiation-shaped .docx by applying "
            "named structural churn transforms to a generated base contract "
            "(issue #565)."
        )
    )
    parser.add_argument("--base", choices=sorted(BASE_DOCUMENT_SPECS), default="mutual-nda")
    parser.add_argument(
        "--transform",
        action="append",
        choices=sorted(TRANSFORMS),
        default=[],
        help="may be repeated; applied in the order given",
    )
    parser.add_argument("--seed", type=int, default=0)
    parser.add_argument("--out", type=Path, help="output .docx path")
    parser.add_argument("--list-bases", action="store_true")
    parser.add_argument("--list-transforms", action="store_true")
    args = parser.parse_args(argv)

    if args.list_bases:
        for name in sorted(BASE_DOCUMENT_SPECS):
            print(name)
        return 0
    if args.list_transforms:
        for name in sorted(TRANSFORMS):
            print(name)
        return 0
    if not args.out:
        parser.error("--out is required unless --list-bases/--list-transforms")

    docx_bytes = build_base_document(args.base)
    for name in args.transform:
        docx_bytes = apply_transform(docx_bytes, name, seed=args.seed)

    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_bytes(docx_bytes)
    print(f"wrote {args.out} ({len(docx_bytes)} bytes)")
    return 0


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())
